import os
import chardet
import csv
from tkinter import Tk, Button, Text, filedialog, Label, Entry
from tkinter import ttk
from passlib.hash import pbkdf2_sha256
from docx import Document


# Update the USERS dictionary with a custom username and password
USERS = {
    "admin": "$pbkdf2-sha256$29000$DcDcDW9LdrwflIO7Zdpa8Q$frccB46zwxKzXwhNqxGBikwESuh2ZfhrR40ziC/fG1g",
    "210804": pbkdf2_sha256.hash("Krister")
}


class CommentingApp:
    def __init__(self):
        self.root = Tk()
        self.root.title("Commenting App")

        # Calculate the center coordinates of the screen
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()
        x = (screen_width - 600) // 2  # Adjust the width as desired
        y = (screen_height - 400) // 2  # Adjust the height as desired

        # Set the window size and position
        self.root.geometry(f"600x400+{x}+{y}")

        # Create a Label for username
        self.username_label = Label(self.root, text="Username:")
        self.username_label.pack()

        # Create an Entry widget for username
        self.username_entry = Entry(self.root)
        self.username_entry.pack()

        # Create a Label for password
        self.password_label = Label(self.root, text="Password:")
        self.password_label.pack()

        # Create an Entry widget for password
        self.password_entry = Entry(self.root, show="*")
        self.password_entry.pack()

        # Create a "Login" button
        self.login_button = Button(self.root, text="Login", command=self.login)
        self.login_button.pack()

        self.comments = {}  # Dictionary to store comments
        self.logged_in = False

    def login(self):
        username = self.username_entry.get()
        password = self.password_entry.get()

        if username in USERS:
            hashed_password = USERS[username]
            if pbkdf2_sha256.verify(password, hashed_password):
                self.logged_in = True
                self.username_label.pack_forget()
                self.username_entry.pack_forget()
                self.password_label.pack_forget()
                self.password_entry.pack_forget()
                self.login_button.pack_forget()
                self.initialize_app()
                print("Login successful!")
            else:
                print("Incorrect password!")
        else:
            print("Username not found!")

    def initialize_app(self):
        # Create an "Open Folder" button
        self.open_button = Button(self.root, text="Open Folder", command=self.open_folder)
        self.open_button.pack()

        # Create a Treeview widget
        self.tree = ttk.Treeview(self.root)
        self.tree.pack(fill="both", expand=True)

        # Configure Treeview columns
        self.tree["columns"] = ("Type")
        self.tree.column("#0", width=300, minwidth=300, stretch="no")
        self.tree.column("Type", width=100, minwidth=100, stretch="no")

        # Configure Treeview column headings
        self.tree.heading("#0", text="File/Folder Name")
        self.tree.heading("Type", text="Type")

        # Bind double-click event to Treeview items
        self.tree.bind("<Double-1>", self.on_item_doubleclick)

    def open_folder(self):
        folder_path = filedialog.askdirectory()
        if folder_path:
            self.populate_treeview(folder_path)

    def populate_treeview(self, folder_path):
        self.tree.delete(*self.tree.get_children())
        items = os.listdir(folder_path)
        for item in items:
            item_path = os.path.join(folder_path, item)
            item_type = "Folder" if os.path.isdir(item_path) else "File"
            self.tree.insert("", "end", text=item, values=(item_type,))

    def on_item_doubleclick(self, event):
        item = self.tree.focus()
        item_text = self.tree.item(item, "text")
        item_type = self.tree.item(item, "values")[0]

        if item_type == "File":
            self.open_file(item_text)

    def open_file(self, file_name):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])

        if file_path:
            try:
                document = Document(file_path)
                text = self.extract_text_from_document(document)
                self.display_text_and_comment(file_name, text)
            except Exception as e:
                print(f"Failed to open file: {file_path}")
                print(e)

    


    def extract_text_from_document(self, document):
        paragraphs = document.paragraphs
        text = ""
        for paragraph in paragraphs:
            text += paragraph.text + "\n"
        return text

    def display_text_and_comment(self, file_name, text):
        if file_name not in self.comments:
            self.comments[file_name] = ""

        if not self.logged_in:
            print("Please log in to add comments.")
            return

        self.root.title(f"Commenting App - {file_name}")

        # Clear the existing widgets
        for widget in self.root.winfo_children():
            widget.pack_forget()

        # Create a Text widget to display the file content
        text_widget = Text(self.root, wrap="word")
        text_widget.insert("1.0", text)
        text_widget.pack(fill="both", expand=True)

        # Create a Label for comment
        comment_label = Label(self.root, text="Comment:")
        comment_label.pack()

        # Create a Text widget for comment
        comment_text = Text(self.root, wrap="word", height=5)
        comment_text.pack()

        # Create a "Add Comment" button
        add_comment_button = Button(self.root, text="Add Comment",
                                    command=lambda: self.add_comment(file_name, comment_text.get("1.0", "end-1c")))
        add_comment_button.pack()

        # Create a "Save" button
        save_button = Button(self.root, text="Save", command=lambda: self.save_comments(file_name))
        save_button.pack()

        # Display the existing comment
        comment_text.insert("1.0", self.comments[file_name])

    def add_comment(self, file_name, comment):
        if file_name in self.comments:
            self.comments[file_name] = comment
            print("Comment added!")
        else:
            print("File not found!")

    def save_comments(self, file_name):
        if file_name in self.comments:
            comment = self.comments[file_name]
            self.comments[file_name] = comment
            self.export_to_word(file_name, comment)
            print("Comments saved!")
        else:
            print("File not found!")

    def export_to_word(self, file_name, comment):
        document = Document()
        document.add_paragraph("Comment:")
        document.add_paragraph(comment)
        output_file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                        initialfile=f"{file_name}_with_comment")
        if output_file_path:
            document.save(output_file_path)


app = CommentingApp()
app.root.mainloop()
